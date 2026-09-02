"""多格式文档解析：把 txt / md / html / pdf / docx / xml / json / csv / tsv / xlsx
统一解析成 Document（元数据一致：title/source/character/url/category/file/aliases）。

为什么需要：知识库原本只认 *.txt（load_knowledge_files 单格式 glob），无法接入
PDF/HTML/DOCX/JSON/CSV/XLSX/MD 等常见史料格式。本模块按扩展名分发，格式差异在
进入向量库前被抹平；load_knowledge_files 改为覆盖全部支持扩展名。

格式约定（与既有 txt 一致，其它格式也支持）：
  - 前置头：# 标题、【来源】【URL】【人物】【分类】+ 来源标注【朝代】【出处】【篇卷】；
    `---` 分隔正文；
  - character 未显式给出时尝试从文件名（biography_张衡_内置.txt）推导，否则留空
    （留空文档仍可被全局检索命中，只是不参与按人物过滤）；
  - doc_type 由 `_内置` 后缀 / 来源=内置知识库 自动判定（persona），显式给出优先；
    朝代缺失时从人物配置补齐；
  - 未知扩展名告警跳过；PDF 无文本（扫描件）明确告警"不支持 OCR"。

离线约束：不依赖 unstructured/pandoc 等重型工具——md 用 TextLoader、html 用
bs4、pdf 用 pypdf、docx 用 python-docx、xlsx 用 pandas/openpyxl，全部已装。
"""
import csv
import json
import re
from pathlib import Path
from typing import Dict, List, Optional

from langchain_core.documents import Document

from src.logger import get_logger

logger = get_logger("document_loader")

SUPPORTED_EXTS = {
    ".txt", ".md", ".markdown", ".html", ".htm", ".pdf", ".docx",
    ".xml", ".json", ".csv", ".tsv", ".xlsx",
}

# 结构化格式（JSON/CSV/XLSX）里允许映射到元数据的列/键白名单。
# 来源标注字段 doc_type/dynasty/book/chapter 由此透传到 _finalize。
_META_KEYS = ("title", "source", "url", "character", "category",
              "doc_type", "dynasty", "book", "chapter")


# ───────────────────────── 元数据与正文抽取 ─────────────────────────

def _extract_body(lines: List[str]) -> str:
    """取 `---` 之后的正文；无 `---` 时跳过开头的标题/元数据/空行取剩余。

    与既有 txt 语义一致：`---` 是"前面都是元数据、后面才是正文"的分隔符。
    """
    for i, ln in enumerate(lines):
        if ln.startswith("---"):
            return "\n".join(lines[i + 1:]).strip()
    for i, ln in enumerate(lines):
        if ln.strip() and not ln.startswith("#") and not ln.startswith("【"):
            return "\n".join(lines[i:]).strip()
    return "\n".join(lines).strip()


def _parse_front_matter(content: str) -> Dict[str, str]:
    """解析前置头元数据（# 标题、【来源】【URL】【人物】【分类】【朝代】【出处】【篇卷】）。"""
    meta: Dict[str, str] = {}
    for line in content.split("\n")[:10]:
        if line.startswith("# "):
            meta["title"] = line[2:].strip()
        elif line.startswith("【来源】"):
            meta["source"] = line[4:].strip()
        elif line.startswith("【URL】"):
            meta["url"] = line[5:].strip()
        elif line.startswith("【人物】"):
            meta["character"] = line[4:].strip()
        elif line.startswith("【分类】"):
            meta["category"] = line[4:].strip()
        elif line.startswith("【朝代】"):
            meta["dynasty"] = line[4:].strip()
        elif line.startswith("【出处】"):
            meta["book"] = line[4:].strip()
        elif line.startswith("【篇卷】"):
            meta["chapter"] = line[4:].strip()
        elif line.startswith("【doc_type】"):
            meta["doc_type"] = line[len("【doc_type】"):].strip()
        elif line.startswith("【白话】"):
            meta["baihua"] = line[4:].strip()
    return meta


def _char_from_filename(stem: str) -> str:
    """从文件名推导人物（约定 biography_<姓名>[_内置]）。"""
    m = re.fullmatch(r"biography_([一-鿿]+)(?:_内置)?", stem)
    return m.group(1) if m else ""


def _infer_doc_type(source: Path, meta: Dict[str, str]) -> str:
    """文档类型：内置生成摘要 → persona（语言风格参考，非事实依据）；否则 → historical。

    判定依据：文件名带 `_内置` 后缀，或来源标记为"内置知识库"。真实史源（ctext/
    维基/公版古籍节选等）不带该后缀 → historical。persona 与 historical 是检索
    双轨（history_agent）与引用标注（block/footer）的基础。
    """
    explicit = (meta.get("doc_type") or "").strip().lower()
    if explicit in ("persona", "historical"):
        return explicit
    if "_内置" in source.stem:
        return "persona"
    if (meta.get("source") or "").strip() == "内置知识库":
        return "persona"
    return "historical"


def _resolve_dynasty(character: str, meta: Dict[str, str]) -> str:
    """朝代：前置头显式给出优先；否则从人物配置补齐（97 人 YAML 均带 dynasty）。"""
    dynasty = (meta.get("dynasty") or "").strip()
    if dynasty:
        return dynasty
    if character:
        try:
            from src.characters import character_manager
            char = character_manager.get_character(character)
            if char:
                return char.dynasty
        except Exception:
            pass
    return ""


def _finalize(source: Path, body: str, meta: Dict[str, str]) -> Optional[Document]:
    """统一包装成 Document：填默认元数据 + 别名 + 来源标注，body 为空返回 None。

    元数据键：title/source/url/character/category/aliases/file 之外，新增
    doc_type（persona|historical）、dynasty（朝代）、book（出处）、chapter（篇卷）。
    """
    if not body.strip():
        logger.warning("文档无正文，跳过: %s", source.name)
        return None
    from src.knowledge.aliases import get_aliases_for

    character = (meta.get("character") or "").strip() or _char_from_filename(source.stem)
    aliases = get_aliases_for(character) if character else []
    # 白话导读（【白话】，gushiwen 抓取产物带）：现代语锚点，嵌入时前置到正文，
    # 缓解"通济渠"vs"大运河"类古今用词代差；标注"白话导读"与原文区分，非事实依据。
    gloss = (meta.get("baihua") or "").strip()
    page_content = f"白话导读：{gloss}\n\n{body}" if gloss else body
    return Document(
        page_content=page_content,
        metadata={
            "title": (meta.get("title") or "").strip() or source.stem,
            "source": (meta.get("source") or "").strip() or "未知",
            "url": (meta.get("url") or "").strip(),
            "character": character,
            "category": (meta.get("category") or "").strip() or "biography",
            "aliases": ",".join(aliases),
            "file": source.name,
            "doc_type": _infer_doc_type(source, meta),
            "dynasty": _resolve_dynasty(character, meta),
            # book 去书名号：front-matter 【出处】《史记》与映射表 book=史记 统一为裸书名，
            # 标签渲染时再统一加《》（避免《《史记》》双重嵌套）。
            "book": (meta.get("book") or "").strip().strip("《》"),
            "chapter": (meta.get("chapter") or "").strip(),
        },
    )


def _load_text_based(source: Path, text: str) -> List[Document]:
    """文本类格式通用路径：前置头 + 正文 → 单个 Document。"""
    meta = _parse_front_matter(text)
    body = _extract_body(text.split("\n"))
    doc = _finalize(source, body, meta)
    return [doc] if doc else []


# ───────────────────────── 各格式解析 ─────────────────────────

def _load_txt(source: Path) -> List[Document]:
    with open(source, "r", encoding="utf-8") as f:
        content = f.read()
    return _load_text_based(source, content)


def _load_md(source: Path) -> List[Document]:
    from langchain_community.document_loaders import TextLoader
    docs = TextLoader(str(source), encoding="utf-8").load()
    return _load_text_based(source, docs[0].page_content if docs else "")


def _load_html(source: Path) -> List[Document]:
    from langchain_community.document_loaders import BSHTMLLoader
    docs = BSHTMLLoader(str(source), open_encoding="utf-8").load()
    return _load_text_based(source, docs[0].page_content if docs else "")


def _load_pdf(source: Path) -> List[Document]:
    from langchain_community.document_loaders import PyPDFLoader
    pages = PyPDFLoader(str(source)).load()
    text = "\n".join(p.page_content for p in pages).strip()
    if not text:
        logger.warning("PDF 无文本层（可能为扫描件），不支持 OCR，跳过: %s", source.name)
        return []
    return _load_text_based(source, text)


def _load_docx(source: Path) -> List[Document]:
    # 不用 Docx2txtLoader：它依赖 docx2txt 包（未装）；python-docx 已装且够用。
    from docx import Document as DocxDocument

    doc = DocxDocument(str(source))
    text = "\n".join(p.text for p in doc.paragraphs)
    return _load_text_based(source, text)


def _load_xml(source: Path) -> List[Document]:
    from bs4 import BeautifulSoup
    with open(source, "r", encoding="utf-8") as f:
        # XML 必须用 XML 解析器（features="xml" → lxml-xml），用默认 HTML 解析器
        # 会把自闭合标签/命名空间/实体按 HTML 语义错误展开，且抛 XMLParsedAsHTMLWarning
        # （如 `<item/>` 被当成 `<item>` 的开标签，后续文本错位）。lxml 已装。
        text = BeautifulSoup(f.read(), "xml").get_text("\n")
    return _load_text_based(source, text)


def _load_json(source: Path) -> List[Document]:
    """JSON 知识格式：数组 [{title, content/text, source, character, url, category}]

    每元素一个 Document（可带各自人物/来源）；非该结构则整体转成单个 Document。
    """
    data = json.loads(source.read_text(encoding="utf-8"))
    if isinstance(data, list) and all(isinstance(d, dict) for d in data):
        out = []
        for d in data:
            body = d.get("content") or d.get("text") or ""
            doc = _finalize(source, str(body).strip(), {str(k): str(v) for k, v in d.items() if k not in ("content", "text")})
            if doc:
                out.append(doc)
        return out
    # 通用兜底：整份 JSON 递归取所有字符串拼成一个 Document
    texts = []

    def _walk(x):
        if isinstance(x, dict):
            for v in x.values():
                _walk(v)
        elif isinstance(x, list):
            for v in x:
                _walk(v)
        elif isinstance(x, str) and x.strip():
            texts.append(x.strip())

    _walk(data)
    return _load_text_based(source, "\n".join(texts))


def _load_csv(source: Path) -> List[Document]:
    """结构化知识 CSV/TSV：表头含 content/text 列时按行生成 Document。

    为什么不用 CSVLoader：它把每行渲染成"列名: 值"的平铺文本，不适合知识库的
    content 语义；用 csv.DictReader 按行解析，能把人物/来源等列直接映射到元数据。
    """
    delimiter = "\t" if source.suffix.lower() == ".tsv" else ","
    with open(source, "r", encoding="utf-8", newline="") as f:
        rows = list(csv.DictReader(f, delimiter=delimiter))
    out = []
    for row in rows:
        row = {k.strip(): (v or "").strip() for k, v in row.items() if k}
        body = row.pop("content", None) or row.pop("text", None) or ""
        if not body:
            continue
        meta = {k: v for k, v in row.items() if k in _META_KEYS}
        doc = _finalize(source, body, meta)
        if doc:
            out.append(doc)
    if not out:
        logger.warning("CSV 无有效行（需 content/text 列）: %s", source.name)
    return out


def _load_xlsx(source: Path) -> List[Document]:
    """Excel：sheet 含 content/text 列时逐行生成 Document，否则整 sheet 拼成一个。"""
    import pandas as pd
    sheets = pd.read_excel(source, sheet_name=None)
    out = []
    for name, df in sheets.items():
        df = df.dropna(how="all")
        if df.empty:
            continue
        cols = [str(c) for c in df.columns]
        if "content" in cols or "text" in cols:
            body_col = "content" if "content" in cols else "text"
            for _, r in df.iterrows():
                body = str(r.get(body_col, "")).strip()
                if not body:
                    continue
                meta = {
                    k: str(r.get(k, "")).strip()
                    for k in _META_KEYS
                    if k in cols
                }
                meta.setdefault("title", f"{name}-{len(out) + 1}")
                doc = _finalize(source, body, meta)
                if doc:
                    out.append(doc)
        else:
            # 整 sheet 拼接为一段（单元格以制表/换行分隔）
            cells = []
            for _, r in df.iterrows():
                cells.append("\t".join(str(v) for v in r if str(v).strip()))
            doc = _finalize(source, "\n".join(cells), {"title": name})
            if doc:
                out.append(doc)
    return out


# ───────────────────────── 统一入口 ─────────────────────────

def load_documents(source) -> List[Document]:
    """按扩展名分发解析文件，返回统一 Document 列表；未知格式返回 []。"""
    path = Path(source)
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTS:
        logger.warning("不支持的文件扩展名 %s，跳过: %s", ext, path.name)
        return []

    loaders = {
        ".txt": _load_txt,
        ".md": _load_md,
        ".markdown": _load_md,
        ".html": _load_html,
        ".htm": _load_html,
        ".pdf": _load_pdf,
        ".docx": _load_docx,
        ".xml": _load_xml,
        ".json": _load_json,
        ".csv": _load_csv,
        ".tsv": _load_csv,
        ".xlsx": _load_xlsx,
    }
    try:
        docs = loaders[ext](path)
    except Exception as e:
        logger.error("解析失败 %s: %s", path.name, e)
        return []
    if docs:
        logger.info("  解析 %s → %d 篇文档", path.name, len(docs))
    return docs
